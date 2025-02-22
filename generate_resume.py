# Copyright (c) 2024 Joseph Mapula. All rights reserved. See LICENSE file for details.
import os
from typing import Optional
from dotenv import load_dotenv
import openai
from openai import OpenAI
import anthropic
from mistralai import Mistral
import PyPDF2
from docx import Document
import argparse
import tempfile
from werkzeug.utils import secure_filename
import hashlib
import shutil
import time
import re
import json
import subprocess 
from typing import Dict
# Import the logging module, which provides a flexible framework for generating log messages in Python
import logging
# Configure the basic settings for the logging system
logging.basicConfig(
    level=logging.DEBUG,  # Set the root logger's level 
    # This means it will capture all logs of severity INFO and above (INFO, WARNING, ERROR, CRITICAL)
    # Logs below this level (like DEBUG) will be ignored unless explicitly set for specific loggers
    
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    # Define the format of log messages:
    # %(asctime)s: Timestamp when the log was created
    # %(name)s: Name of the logger (in this case, it will be '__main__' unless specified otherwise)
    # %(levelname)s: Severity level of the log (e.g., INFO, WARNING, ERROR)
    # %(message)s: The actual log message
    
    datefmt='%Y-%m-%d %H:%M:%S'
    # Specify the date/time format in the log messages
)

# Create a logger instance for this module
logger = logging.getLogger(__name__)
# __name__ is a special variable in Python that holds the name of the current module
# In the main script, __name__ is set to '__main__'
# This allows you to have separate loggers for different modules in larger applications

# Now you can use logger to create log messages, for example:
# logger.debug("This is a debug message")
# logger.info("This is an info message")
# logger.warning("This is a warning message")
# logger.error("This is an error message")
# logger.critical("This is a critical message")

# The logging setup above ensures that:
# 1. All log messages are consistently formatted
# 2. Each log includes a timestamp, logger name, log level, and the message
# 3. Only logs of INFO level and above will be displayed/saved (due to level=logging.INFO)
# 4. You can easily adjust the logging level and format globally by modifying the basicConfig

# Load environment variables from the .env file
# This allows us to securely store and access API keys without hardcoding them
load_dotenv()

# Set up API keys for OpenAIm, Anthropic, and Mistral (missing will return none)
openai_api_key = os.getenv("OPENAI_API_KEY")
anthropic_api_key = os.getenv("ANTHROPIC_API_KEY")
mistral_api_key = os.getenv("MISTRAL_API_KEY")

def validate_api_keys():
    if not openai_api_key:
        logger.error("OpenAI API key is missing. Please check your .env file.")
        raise ValueError("OpenAI API key is required")
    if not anthropic_api_key:
        logger.error("Anthropic API key is missing. Please check your .env file.")
        raise ValueError("Anthropic API key is required")
    if not mistral_api_key: 
        logger.error("Mistral API key is missing. Please check your .env file.")
        raise ValueError("Mistral API key is required")
    logger.info("API keys validated successfully.")
# Validate API keys before proceeding
validate_api_keys()

# Create clients for each provider
try:
    anthropic_client = anthropic.Anthropic(api_key=anthropic_api_key)
    openai_client = OpenAI(api_key=openai_api_key)
    mistral_client = Mistral(api_key=mistral_api_key)
    logger.info("Clients created successfully.")
except Exception as e:
    logger.error(f"Error creating API clients: {e}")
    raise

# Define folders
UPLOAD_FOLDER = 'uploads'
TEMP_FOLDER = 'temp'
OUTPUT_FOLDER = 'outputs'
LOG_FOLDER = 'logs'

# Ensure necessary folders exist
for folder in [UPLOAD_FOLDER, TEMP_FOLDER, OUTPUT_FOLDER, LOG_FOLDER]:
    if not os.path.exists(folder):
        os.makedirs(folder)
        logging.info(f"Created folder: {folder}")


# Define Knowledge Base storage file
KNOWLEDGE_BASE_FILE = "knowledge_base.json"

# Ensure Knowledge Base exists
def initialize_knowledge_base():
    if not os.path.exists(KNOWLEDGE_BASE_FILE):
        with open(KNOWLEDGE_BASE_FILE, "w") as f:
            json.dump({}, f)
        logging.info("Initialized new knowledge base.")

# Read and parse different file types
def parse_resume(file_path: str) -> Dict:
    """
    Extracts structured data from a resume (PDF, DOCX, TXT, MD) and stores it in a structured format.
    """
    _, file_extension = os.path.splitext(file_path)
    content = ""
    
    try:
        if file_extension.lower() == ".pdf":
            with open(file_path, "rb") as file:
                reader = PyPDF2.PdfReader(file)
                content = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
        elif file_extension.lower() in [".docx", ".doc"]:
            doc = Document(file_path)
            content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        elif file_extension.lower() in [".txt", ".md"]:
            with open(file_path, "r", encoding="utf-8") as file:
                content = file.read()
        else:
            logging.error(f"Unsupported file format: {file_extension}")
            return {}
    except Exception as e:
        logging.error(f"Error reading file {file_path}: {e}")
        return {}
    
    return extract_structured_data(content)

# Extract structured data from raw text
def extract_structured_data(text: str) -> Dict:
    """
    Basic structured parsing of resume text.
    """
    data = {
        "name": "Unknown",
        "email": "Unknown",
        "phone": "Unknown",
        "linkedin": "Unknown",
        "work_experience": [],
        "education": [],
        "skills": []
    }
    
    lines = text.split("\n")
    for line in lines:
        if "@" in line and ".com" in line:
            data["email"] = line.strip()
        elif "linkedin.com" in line:
            data["linkedin"] = line.strip()
        elif any(keyword in line.lower() for keyword in ["work experience", "professional experience"]):
            data["work_experience"].append(line.strip())
        elif "education" in line.lower():
            data["education"].append(line.strip())
        elif any(keyword in line.lower() for keyword in ["skills", "expertise"]):
            data["skills"].append(line.strip())
    
    return data

# Generate a unique user ID based on extracted data
def generate_user_id(name: str, email: str) -> str:
    if email != "Unknown":
        return hashlib.md5(email.encode()).hexdigest()  # Unique but anonymized ID
    return f"user_{int(time.time())}"  # Fallback to timestamp-based ID

# Save structured resume data to Knowledge Base
def save_to_knowledge_base(user_id: str, resume_data: Dict):
    """
    Saves parsed resume data to the knowledge base under a unique user ID.
    """
    initialize_knowledge_base()
    
    try:
        with open(KNOWLEDGE_BASE_FILE, "r") as f:
            knowledge_base = json.load(f)
        
        knowledge_base[user_id] = resume_data
        
        with open(KNOWLEDGE_BASE_FILE, "w") as f:
            json.dump(knowledge_base, f, indent=4)
        logging.info(f"Resume data saved for user {user_id}.")
    except Exception as e:
        logging.error(f"Error saving to knowledge base: {e}")

# Main function to process resume
def process_resume(file_path: str):
    """
    Parses and saves a resume to the knowledge base.
    """
    resume_data = parse_resume(file_path)
    if resume_data:
        user_id = generate_user_id(resume_data.get("name", "Unknown"), resume_data.get("email", "Unknown"))
        save_to_knowledge_base(user_id, resume_data)
    else:
        logging.error("Failed to extract structured data from resume.")

def generate_unique_filename(job_title, company, model, extension=".md"):
    """Generates a unique filename using job title, company, model, and timestamp."""
    timestamp = time.strftime("%Y-%m-%d_%H%M%S")

    safe_job_title = secure_filename((job_title or "Unknown_Job_Title").replace(" ", "_").lower())
    safe_company = secure_filename((company or "Unknown_Company").replace(" ", "_").lower())

    filename = f"resume_{safe_company}_{safe_job_title}_{model}_{timestamp}{extension}"
    
    return os.path.join(OUTPUT_FOLDER, filename)


def extract_job_title_and_company_regex(text):
    title_pattern = re.search(r'(?i)(?:title|position):\s*(.*)', text)
    company_pattern = re.search(r'(?i)(?:company|employer|organization):\s*(.*)', text)

    job_title = title_pattern.group(1).strip() if title_pattern else "Unknown_Job_Title"
    company_name = company_pattern.group(1).strip() if company_pattern else "Unknown_Company"

    return job_title, company_name


def extract_job_details_with_llm(text):
    """Use an LLM to extract the job title and company from a description."""
    prompt = f"""
    Extract the job title and company name from the following job description:
    
    {text}

    Respond in JSON format:
    {{"job_title": "<job title>", "company": "<company name>"}}
    """

    response = mistral_client.chat.completions.create(
        model="mistral-large-latest",  
        messages=[{"role": "user", "content": prompt}],
        max_tokens=100,
        temperature=0  # Reduce creativity for precise extraction
    )

    result = response.choices[0].message.content.strip()

    # Safely parse JSON response
    try:
        parsed_result = json.loads(result)
        return parsed_result
    except json.JSONDecodeError:
        logging.error(f"Failed to parse LLM response: {result}")
        return {"job_title": "Unknown_Job_Title", "company": "Unknown_Company"}

def get_job_title_and_company(text):
    job_title, company_name = extract_job_title_and_company_regex(text)

    if not job_title or not company_name:
        logging.info("Falling back to LLM for job details extraction.")
        llm_result = extract_job_details_with_llm(text)
        job_title, company_name = llm_result.get("job_title"), llm_result.get("company")

    return job_title, company_name or "Unknown Company"


def cleanup_temp_folder():
    """Deletes temporary files after processing."""
    try:
        shutil.rmtree(TEMP_FOLDER)
        os.makedirs(TEMP_FOLDER)  # Recreate empty temp folder
        logging.info("Temporary folder cleaned up.")
    except Exception as e:
        logging.error(f"Error cleaning temp folder: {e}")

        
def read_file(file_path):
    """
    Read content from various file types.
    
    Args:
        file_path (str): Path to the file.
    
    Returns:
        str: Content of the file.
    """
    _, file_extension = os.path.splitext(file_path)
    
    try:
        if file_extension.lower() == '.pdf':
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfFileReader(file)
                content = ""
                for page_num in range(reader.numPages):
                    content += reader.getPage(page_num).extract_text()
            logger.info(f"Successfully read PDF file: {file_path}")
            return content
        elif file_extension.lower() in ['.docx', '.doc']:
            doc = Document(file_path)
            content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            logger.info(f"Successfully read Word document: {file_path}")
            return content
        elif file_extension.lower() in ['.txt', '.md']:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            logger.info(f"Successfully read file: {file_path}")
            return content
        else:
            logger.error(f"Unsupported file type: {file_extension}")
            return ""
    except Exception as e:
        logger.error(f"Error reading file {file_path}: {str(e)}")
        return ""

def generate_resume(
    job_description: str,
    background_info: str,
    best_practices: str,
    provider: str = "mistral",
    model: str = "mistral-small-latest",
    custom_prompt: Optional[str] = None,
    max_tokens: int = 4000
) -> str:    
    """
    Generate a tailored resume based on the provided information.

    This function takes various inputs and uses an AI model to generate a resume.
    It supports multiple AI providers and models, and allows for a custom prompt.

    Args:
        job_description (str): The description of the job being applied for.
        background_info (str): The applicant's professional background.
        best_practices (str): Best practices for resume writing.
        provider (str): The AI provider to use ('openai', 'anthropic', or 'mistral').
        model (str): The specific model to use (e.g., 'gpt-3.5-turbo' for OpenAI).
        custom_prompt (str, optional): A custom prompt to override the default.
        max_tokens (int): The maximum number of tokens in the generated resume.
                          Default is 1000, which is about 750 words or 1.5 pages.
                          Adjust this value based on desired resume length.

    Returns:
        str: The generated resume content.

    Raises:
        ValueError: If an unsupported provider is specified.
    """

    # Define the default prompt structure
    default_prompt = f"""
    Job Description:
    {job_description}
    -----
    My Professional Background:
    {background_info}
    -----
    #Best Practices for Job Applications:
    {best_practices}
    -----
    Based on the job description, my professional background, and the best practices for job applications, 
    generate a tailored resume in markdown format. The resume should highlight relevant skills and experiences 
    that match the job requirements. 

    Structure it like:
    # Name
    ## Contact
    ## Experience
    ## Education
    ## Skills

    Ensure that the content is concise, impactful, and directly relevant to the product manager position described.
    Generate a professional resume in strict Markdown only. 
    Very important: No extraneous text. Your output will be converted to a pdf as is and needs to only include the final product. 
    Please keep in mind the differences between the job description, my professional background, and the best practices provided. 
    Be very careful not to mix up the information in the background info with that of the best practices. 
    All information should be truthful based on my professional background info provided, but feel free to use your best judgement to 
    rephrase (within reason) to match the job description and underlying desires of the recruiter, hiring manager, and team. 
    Think about the user experience of all those who will read the results as well as any ATS. 
    Make efficient use of space keeping things on the same line when possible. 
    For example, by grouping relevant skills
    Structure it like:
    # Name
    ## Contact
    ## Experience
    ## Education
    ## Skills
    -----
    Wait. Understand. Reflect. Be thoughtful, thorough, detailed, and organized. Do your best. 
    """

    # Use the custom prompt if provided, otherwise use the default
    prompt = custom_prompt if custom_prompt else default_prompt

    logger.info(f"Generating resume using {provider} ({model})")
    logger.debug(f"Prompt preview: {prompt[:200]}...")
    # Use the appropriate AI provider based on the 'provider' parameter
    if provider == "openai":
        # Create a chat completion using OpenAI's API
        response = openai_client.chat.completions.create(
            model=model,  # The specific model to use (e.g., 'gpt-3.5-turbo')
            messages=[
                {"role": "system", "content": "You are an expert resume writer specializing in product management positions. You deliver results in a ready-to-go format with no other output."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=max_tokens,  # Limit the response length. Adjust this for longer or shorter resumes.
            n=1,  # Generate only one completion
            stop=None,  # No custom stop sequence
            temperature=0.7,  # Control randomness (0.7 is moderately creative)
        )
        # Extract and return the generated content from the API response
        return response.choices[0].message.content.strip()
    elif provider == "anthropic":
        # Create a completion using Anthropic's API
        message = anthropic_client.messages.create(
            model=model,  # The specific Claude model to use
            max_tokens=max_tokens,  # Limit the response length. Adjust this for longer or shorter resumes.
            messages=[
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            temperature=0.7,  # Control randomness (0.7 is moderately creative)
        )
        # Extract and return the generated content from the API response
        return message.content[0].text.strip()
    elif provider == "mistral":
        # Create a chat completion using Mistral's API
        response = mistral_client.chat.complete(
            model=model,
            messages=[{
                "role": "user",
                "content": prompt
            }],
            temperature=0.7,  # Control randomness (0.7 is moderately creative)
            max_tokens=max_tokens
        )

        # Extract and return the generated content from the API response
        return response.choices[0].message.content.strip()
    else:
        # Raise an error if an unsupported provider is specified
        raise ValueError("Unsupported provider. Please use 'openai', 'anthropic', or 'mistral'.")
# Temperature spectrum:
# 0.0: Deterministic/repetitive. Always generates the most probable completion.
# 0.3: Conservative/focused. Generates more predictable and coherent text.
# 0.7: Moderately creative. Balances coherence with novelty. (Our current setting)
# 1.0: Very creative. More diverse outputs, but may occasionally be less coherent.
# 2.0: Highly unpredictable. Can generate more unusual or creative outputs, but with higher risk of incoherence.


def process_uploaded_files(job_description_path, background_info_path, best_practices_path):
    """
    Reads content from file paths safely.

    Args:
        job_description_path (str): Path to job description file.
        background_info_path (str): Path to background information file.
        best_practices_path (str): Path to best practices file.

    Returns:
        tuple: Contents of job description, background info, and best practices.
    """

    def read_file_safe(file_path):
        """Reads a file safely without modifying it."""
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return ""

        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error reading file {file_path}: {str(e)}")
            return ""

    job_content = read_file_safe(job_description_path)
    background_content = read_file_safe(background_info_path)
    practices_content = read_file_safe(best_practices_path)

    return job_content, background_content, practices_content

def save_resume(resume_content, job_title, company, model, format="md"):
    """Saves the generated resume in the appropriate format."""
    format_mapping = {
        "md": ".md",
        "docx": ".docx",
        "pdf": ".pdf"
    }
    
    extension = format_mapping.get(format, ".md")
    output_path = generate_unique_filename(job_title, company, model, extension)
    
    try:
        # Remove any wrapping backticks before saving
        cleaned_content = resume_content.strip('`').strip()
        
        with open(output_path, "w", encoding="utf-8") as file:
            file.write(cleaned_content)
        logging.info(f"Resume saved successfully: {output_path}")
        return output_path
    except Exception as e:
        logging.error(f"Error saving resume ({output_path}): {e}")
        return None




def clean_markdown_file(md_path: str) -> str:
    """
    Cleans the Markdown file by removing problematic starting characters or extra spaces.
    
    Args:
        md_path (str): Path to the Markdown file.

    Returns:
        str: Path to the cleaned Markdown file.
    """
    try:
        with open(md_path, 'r', encoding='utf-8') as file:
            lines = file.readlines()

        # Remove leading backticks or code block starts
        cleaned_lines = []
        for line in lines:
            # Remove code block fences or accidental ```markdown
            if line.strip().startswith("```"):
                continue
            cleaned_lines.append(line)

        # Overwrite the file with cleaned content
        with open(md_path, 'w', encoding='utf-8') as file:
            file.writelines(cleaned_lines)

        logging.info(f"Markdown file cleaned: {md_path}")
        return md_path

    except Exception as e:
        logging.error(f"Error cleaning Markdown file: {e}")
        return md_path  # Return original if cleaning fails

def convert_md_to_pdf(md_path: str, engine: str = "xelatex") -> str:
    """Convert Markdown to PDF using Pandoc with custom layout settings."""
    # Clean the markdown before conversion
    md_path = clean_markdown_file(md_path)

    pdf_path = md_path.replace(".md", f"_{engine}.pdf")
    try:
        result = subprocess.run([
            "pandoc", md_path, "-o", pdf_path,
            "--pdf-engine", engine,
            "--variable=geometry:margin=1in",  # Apply 1-inch margins
            "--wrap=auto"                      # Ensure lines wrap properly
        ], capture_output=True, text=True, check=True)
        logging.info(f"Converted to PDF using {engine}: {pdf_path}")
    except subprocess.CalledProcessError as e:
        logging.error(f"Pandoc PDF conversion with {engine} failed: {e.stderr}")
        return ""
    return pdf_path

def convert_md_to_docx(md_path: str, template_path: str) -> str:
    """Convert Markdown to DOCX using Pandoc with a Word template."""
    # Clean Markdown before conversion
    md_path = clean_markdown_file(md_path)

    docx_path = md_path.replace(".md", ".docx")
    try:
        result = subprocess.run([
            "pandoc", md_path, "-o", docx_path,
            "--reference-doc", template_path,
            "--from=markdown", "--standalone"
        ], capture_output=True, text=True, check=True)
        logging.info(f"Converted to DOCX: {docx_path}")
    except subprocess.CalledProcessError as e:
        logging.error(f"Pandoc DOCX conversion failed: {e.stderr}")
        return ""
    return docx_path


# def test_all_conversions(md_path: str, docx_template_path: str):
#     """Run DOCX and multiple PDF engine conversions for testing."""
#     logging.info("Starting conversion tests...")
#     md_path = clean_markdown_file(md_path)
#     # DOCX Conversion
#     docx_output = convert_md_to_docx(md_path, docx_template_path)

#     # PDF Conversions with different engines
#     pdf_engines = ["pdflatex", "xelatex", "lualatex"]
#     pdf_outputs = {}
#     for engine in pdf_engines:
#         pdf_outputs[engine] = convert_md_to_pdf(md_path, engine=engine)

#     logging.info("Conversion tests completed.")
#     return {
#         "docx": docx_output,
#         "pdf": pdf_outputs  }

# Example usage
if __name__ == "__main__":
    md_file = r"C:\Users\josep\Documents\programming_projects\GitHub\ResGen-Private\outputs\resume_unknown_company_108800-_mistral-large-latest_2025-02-21_152125.md"  # Replace with actual path
    docx_template = r"C:\Users\josep\Documents\programming_projects\GitHub\ResGen-Private\uploads\pandoc_resume_template.docx"  # Replace with actual template path

    results = test_all_conversions(md_file, docx_template)
    logging.info(f"Results: {results}")


def main(job_description_path, background_info_path, best_practices_path, selected_models=None, log_level=logging.INFO, generate_docx=True, generate_pdf=False):
    """
    Runs resume generation for the selected AI models and converts to DOCX/PDF if specified.

    Args:
        job_description_path (str): Path to the job description file.
        background_info_path (str): Path to the background information file.
        best_practices_path (str): Path to the best practices file.
        selected_models (list, optional): List of models to run (e.g., ["openai", "mistral"]).
        log_level (int, optional): Logging level.
        generate_docx (bool, optional): Whether to generate DOCX output.
        generate_pdf (bool, optional): Whether to generate PDF output.
    """
    logging.getLogger().setLevel(log_level)
    template_path = r"C:\Users\josep\Documents\programming_projects\GitHub\ResGen-Private\uploads\pandoc_resume_template.docx"

    # Default: Run all models if none are selected
    if selected_models is None:
        selected_models = ["mistral"]

    try:
        job_description, background_info, best_practices = process_uploaded_files(
            job_description_path, background_info_path, best_practices_path
        )

        job_title, company_name = extract_job_title_and_company_regex(job_description)

        for model_provider, model_name in [
            ("openai", "gpt-4o-mini"),
            ("anthropic", "claude-3-5-haiku-latest"),
            ("mistral", "mistral-large-latest")
        ]:
            if model_provider in selected_models:
                try:
                    resume_md = generate_resume(
                        job_description, background_info, best_practices, provider=model_provider, model=model_name
                    )
                    # Save Markdown
                    md_path = save_resume(resume_md, job_title, company_name, model_name, format="md")
                    
                    # Clean Markdown before conversion
                    md_path = clean_markdown_file(md_path)

                    # Convert to DOCX if enabled
                    if generate_docx:
                        convert_md_to_docx(md_path, template_path)


                    # Convert to PDF if enabled
                    if generate_pdf:
                        convert_md_to_pdf(md_path)

                except Exception as e:
                    logging.error(f"Failed to process {model_provider} ({model_name}): {e}")

        logging.info("Resumes generated and converted successfully!")

    except FileNotFoundError as e:
        logging.error(f"File not found: {str(e)}")
    except ValueError as e:
        logging.error(f"Invalid input: {str(e)}")
    except IOError as e:
        logging.error(f"I/O error occurred: {str(e)}")
    except anthropic.APIError as e:
        logging.error(f"Anthropic API error: {str(e)}")
    except openai.APIError as e:
        logging.error(f"OpenAI API error: {str(e)}")
    except openai.APIConnectionError as e:
        logging.error(f"Failed to connect to OpenAI API: {e}")
    except openai.RateLimitError as e:
        logging.error(f"OpenAI API request exceeded rate limit: {e}")
    except Exception as e:
        logging.error(f"An unexpected error occurred: {str(e)}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate resumes using AI")
    parser.add_argument("job_description", help="Path to the job description file")
    parser.add_argument("background_info", help="Path to the background information file")
    parser.add_argument("best_practices", help="Path to the best practices file")
    parser.add_argument("--models", nargs="+", choices=['openai', 'anthropic', 'mistral'], 
                        help="Specify which models to run (default: all)")
    parser.add_argument("--log-level", choices=['DEBUG', 'INFO', 'WARNING', 'ERROR', 'CRITICAL'],
                        default='INFO', help="Set the logging level")
    parser.add_argument("--generate-docx", action="store_true", 
                        help="Generate DOCX output in addition to Markdown")
    parser.add_argument("--generate-pdf", action="store_true", 
                        help="Generate PDF output in addition to Markdown")

    args = parser.parse_args()
    log_level = getattr(logging, args.log_level)
    
    # Pass selected inputs to main()
    main(args.job_description, args.background_info, args.best_practices, selected_models=args.models, 
         log_level=log_level, generate_docx=args.generate_docx, generate_pdf=args.generate_pdf)

    # TODO: Create a user interface for inputting information and selecting options
    # TODO: Add a feedback mechanism for iterating on the generated resumes

